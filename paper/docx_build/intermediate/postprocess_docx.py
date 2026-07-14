from __future__ import annotations

from hashlib import sha256
import json
from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[3]
INPUT = ROOT / "paper/docx_build/intermediate/DISSERTATION_PANDOC_RAW.docx"
OUTPUT = ROOT / "paper/docx_build/DISSERTATION_NTU_TEMPLATE_DRAFT.docx"
AUDIT = ROOT / "paper/docx_build/reports/POSTPROCESS_AUDIT.json"

TITLE = "A Canonical Framework for Accuracy–Explainability Trade-off Analysis in Native-Path Knowledge Graph Recommenders"
FRONT_HEADINGS = {
    "Declaration",
    "Acknowledgements",
    "Table of Contents",
    "List of Figures",
    "List of Tables",
    "Abstract",
}
MAJOR_HEADINGS = FRONT_HEADINGS | {"References", "Appendices"}
FIELD_PLACEHOLDERS = {
    "[Table of Contents to be updated in Microsoft Word before final PDF export.]": (
        'TOC \\o "1-3" \\h \\z \\u',
        "[Table of Contents to be updated in Microsoft Word before final PDF export.]",
    ),
    "[List of Figures to be updated in Microsoft Word before final PDF export.]": (
        "TOC \\f F \\h \\z",
        "[List of Figures to be updated in Microsoft Word before final PDF export.]",
    ),
    "[List of Tables to be updated in Microsoft Word before final PDF export.]": (
        "TOC \\f T \\h \\z",
        "[List of Tables to be updated in Microsoft Word before final PDF export.]",
    ),
}


def digest(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_table_width_and_layout(table, column_widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(column_widths)))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
    if len(grid_columns) != len(column_widths):
        raise RuntimeError("Table grid-column count does not match logical columns")
    for grid_column, width in zip(grid_columns, column_widths):
        grid_column.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, column_widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def allocate_column_widths(table, total_width_twips: int) -> list[int]:
    scores = []
    for column in table.columns:
        texts = [cell.text.strip() for cell in column.cells]
        max_characters = max((len(text) for text in texts), default=1)
        max_word = max((len(word) for text in texts for word in re.findall(r"\S+", text)), default=1)
        score = min(max_word, 34) + min(max_characters ** 0.5, 16)
        scores.append(max(score, 8))
    minimum = 720 if len(scores) >= 7 else 820 if len(scores) >= 5 else 980
    minimum_total = minimum * len(scores)
    if minimum_total >= total_width_twips:
        minimum = max(500, total_width_twips // len(scores) - 10)
        minimum_total = minimum * len(scores)
    remaining = total_width_twips - minimum_total
    score_total = sum(scores)
    widths = [minimum + int(remaining * score / score_total) for score in scores]
    widths[-1] += total_width_twips - sum(widths)
    return widths


def section_properties(base, landscape: bool) -> OxmlElement:
    properties = deepcopy(base)
    section_type = properties.find(qn("w:type"))
    if section_type is None:
        section_type = OxmlElement("w:type")
        properties.insert(0, section_type)
    section_type.set(qn("w:val"), "nextPage")
    page_size = properties.find(qn("w:pgSz"))
    if page_size is None:
        raise RuntimeError("Template section has no page-size definition")
    width = page_size.get(qn("w:w"))
    height = page_size.get(qn("w:h"))
    if landscape:
        page_size.set(qn("w:w"), height)
        page_size.set(qn("w:h"), width)
        page_size.set(qn("w:orient"), "landscape")
    else:
        page_size.set(qn("w:w"), width)
        page_size.set(qn("w:h"), height)
        page_size.attrib.pop(qn("w:orient"), None)
    return properties


def make_section_break(properties: OxmlElement) -> OxmlElement:
    paragraph = OxmlElement("w:p")
    paragraph_properties = OxmlElement("w:pPr")
    paragraph_properties.append(properties)
    paragraph.append(paragraph_properties)
    return paragraph


def find_table_caption(table) -> Paragraph:
    node = table._tbl.getnext()
    while node is not None:
        if node.tag == qn("w:tbl"):
            break
        if node.tag == qn("w:p"):
            paragraph = Paragraph(node, table._parent)
            if paragraph.text.strip().startswith("Table "):
                return paragraph
            if paragraph.text.strip():
                break
        node = node.getnext()
    raise RuntimeError("Could not find the caption immediately following a table")


def clear_paragraph_content(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def append_complex_field(paragraph, instruction: str, display_text: str) -> None:
    clear_paragraph_content(paragraph)
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run.append(begin)
    paragraph._p.append(begin_run)

    instruction_run = OxmlElement("w:r")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = f" {instruction} "
    instruction_run.append(instruction_text)
    paragraph._p.append(instruction_run)

    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    paragraph._p.append(separate_run)

    result_run = OxmlElement("w:r")
    result_props = OxmlElement("w:rPr")
    italic = OxmlElement("w:i")
    result_props.append(italic)
    result_run.append(result_props)
    result_text = OxmlElement("w:t")
    result_text.text = display_text
    result_run.append(result_text)
    paragraph._p.append(result_run)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph._p.append(end_run)


def append_tc_field(paragraph, caption_text: str, identifier: str) -> None:
    safe_text = caption_text.replace('"', "'")
    run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run.append(begin)
    paragraph._p.append(run)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    vanish = OxmlElement("w:vanish")
    run_props.append(vanish)
    run.append(run_props)
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f' TC "{safe_text}" \\f {identifier} \\l 1 '
    run.append(instruction)
    paragraph._p.append(run)

    run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run.append(end)
    paragraph._p.append(run)


def set_update_fields(document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


document = Document(INPUT)
if not document.paragraphs or not document.tables or not document.inline_shapes:
    raise RuntimeError("Pandoc DOCX is structurally incomplete")

if "Caption" not in [style.name for style in document.styles]:
    caption_style = document.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
else:
    caption_style = document.styles["Caption"]
caption_style.font.name = "Times New Roman"
caption_style.font.size = Pt(10)
caption_style.font.italic = False
caption_style.paragraph_format.space_before = Pt(3)
caption_style.paragraph_format.space_after = Pt(8)
caption_style.paragraph_format.keep_together = True

normal_style = document.styles["Normal"]
normal_style.font.name = "Times New Roman"
normal_style.font.size = Pt(11)

document.core_properties.title = TITLE
document.core_properties.author = ""
document.core_properties.subject = "NTU dissertation draft for manual inspection"
document.core_properties.keywords = "knowledge graph recommender systems; explainability; dissertation"

front_heading_count = 0
chapter_heading_count = 0
section_heading_count = 0
field_count = 0
figure_caption_count = 0
table_caption_count = 0
image_paragraph_count = 0

for index, paragraph in enumerate(document.paragraphs):
    text = paragraph.text.strip()
    if index == 0 and text == TITLE:
        paragraph.style = document.styles["Title"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(18)
        continue

    if text in FRONT_HEADINGS:
        paragraph.style = document.styles["Heading 1"]
        paragraph.paragraph_format.page_break_before = True
        paragraph.paragraph_format.keep_with_next = True
        front_heading_count += 1
    elif paragraph.style.name == "Heading 1" and re.match(r"^Chapter [1-6]\s+\S", text):
        paragraph.style = document.styles["Heading 1"]
        paragraph.paragraph_format.page_break_before = True
        paragraph.paragraph_format.keep_with_next = True
        chapter_heading_count += 1
    elif text in {"References", "Appendices"}:
        paragraph.style = document.styles["Heading 1"]
        paragraph.paragraph_format.page_break_before = True
        paragraph.paragraph_format.keep_with_next = True
    elif paragraph.style.name == "Heading 2" and re.match(r"^(?:[1-6]\.[0-9]+|Appendix [A-F]:)", text):
        paragraph.style = document.styles["Heading 2"]
        paragraph.paragraph_format.keep_with_next = True
        section_heading_count += 1

    if text in FIELD_PLACEHOLDERS:
        instruction, display = FIELD_PLACEHOLDERS[text]
        append_complex_field(paragraph, instruction, display)
        paragraph.style = document.styles["Normal"]
        paragraph.paragraph_format.space_after = Pt(12)
        field_count += 1
        continue

    caption_match = re.match(r"^(Figure|Table)\s+((?:\d+(?:\.\d+)*)|(?:[A-Z]\.\d+))\.\s+", text)
    if caption_match:
        kind = caption_match.group(1)
        paragraph.style = caption_style
        paragraph.paragraph_format.keep_together = True
        if kind == "Figure":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            append_tc_field(paragraph, text, "F")
            figure_caption_count += 1
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            append_tc_field(paragraph, text, "T")
            table_caption_count += 1

    if paragraph._p.xpath(".//w:drawing"):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.keep_with_next = True
        image_paragraph_count += 1
    elif paragraph.style.name == "Normal" and text and not text.startswith("["):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.widow_control = True

for section in document.sections:
    usable_width = section.page_width - section.left_margin - section.right_margin
    max_image_width = int(usable_width)
    for shape in document.inline_shapes:
        if shape.width > max_image_width:
            ratio = max_image_width / shape.width
            shape.width = max_image_width
            shape.height = int(shape.height * ratio)

base_section = document.sections[-1]
base_section_properties = document._element.body.sectPr
portrait_table_width = int(
    (base_section.page_width - base_section.left_margin - base_section.right_margin) / 635 * 0.96
)
landscape_table_width = int(
    (base_section.page_height - base_section.left_margin - base_section.right_margin) / 635 * 0.96
)
landscape_table_numbers = []
wide_table_numbers = []

for table_number, table in enumerate(document.tables, 1):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    column_count = len(table.columns)
    max_cell_characters = max((len(cell.text) for row in table.rows for cell in row.cells), default=0)
    wide_table = column_count >= 5 or (column_count == 4 and max_cell_characters > 110)
    if wide_table:
        wide_table_numbers.append(table_number)
    target_width = portrait_table_width
    set_table_width_and_layout(table, allocate_column_widths(table, target_width))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    if table.rows:
        set_repeat_table_header(table.rows[0])
    font_size = 7 if column_count >= 8 else 7.5 if column_count == 7 else 8 if column_count == 6 else 8.5 if column_count == 5 else 9
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(font_size)
                    if row_index == 0:
                        run.font.bold = True

set_update_fields(document)
document.save(OUTPUT)

reopened = Document(OUTPUT)
max_available_width = min(
    section.page_width - section.left_margin - section.right_margin for section in reopened.sections
)
oversize_images = [
    index + 1
    for index, shape in enumerate(reopened.inline_shapes)
    if shape.width > max_available_width
]
heading_texts = [
    paragraph.text.strip()
    for paragraph in reopened.paragraphs
    if paragraph.style.name.startswith("Heading")
]
missing_chapters = [
    f"Chapter {chapter}" for chapter in range(1, 7) if not any(text.startswith(f"Chapter {chapter}") for text in heading_texts)
]
all_text = "\n".join(paragraph.text for paragraph in reopened.paragraphs)

AUDIT.write_text(
    json.dumps(
        {
            "input": str(INPUT.relative_to(ROOT)),
            "input_sha256": digest(INPUT),
            "output": str(OUTPUT.relative_to(ROOT)),
            "output_sha256": digest(OUTPUT),
            "paragraphs": len(reopened.paragraphs),
            "tables": len(reopened.tables),
            "inline_shapes": len(reopened.inline_shapes),
            "image_relationships": sum(
                1 for relationship in reopened.part.rels.values() if "image" in relationship.reltype
            ),
            "front_headings_promoted": front_heading_count,
            "chapter_headings": chapter_heading_count,
            "section_headings_checked": section_heading_count,
            "toc_lof_lot_fields": field_count,
            "figure_captions": figure_caption_count,
            "table_captions": table_caption_count,
            "image_paragraphs": image_paragraph_count,
            "oversize_images": oversize_images,
            "missing_chapters": missing_chapters,
            "references_present": "References" in heading_texts,
            "appendices_present": "Appendices" in heading_texts,
            "manual_placeholders_preserved": all(
                marker in all_text
                for marker in [
                    "[Name to be filled]",
                    "[Supervisor to be filled]",
                    "[Date to be filled]",
                    "[Declaration text to be filled according to the NTU template.]",
                ]
            ),
            "section_count": len(reopened.sections),
            "landscape_table_numbers": landscape_table_numbers,
            "wide_table_numbers_requiring_render_check": wide_table_numbers,
            "landscape_section_count": sum(
                1 for section in reopened.sections if section.orientation is not None and section.page_width > section.page_height
            ),
            "low_level_ooxml_reasons": [
                "TOC/List-of-Figures/List-of-Tables field insertion",
                "TC field registration for existing figure and table captions",
                "Word updateFields setting",
                "table width, autofit, repeated-header, and row-split controls",
            ],
        },
        indent=2,
    )
    + "\n",
    encoding="utf-8",
)
