from __future__ import annotations

from collections import Counter
from hashlib import sha256
import json
from pathlib import Path
import re
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[3]
BUILD = ROOT / "paper/docx_build"
SOURCE = ROOT / "paper/current_dissertation/FULL_DISSERTATION_CURRENT.md"
PREPARED = BUILD / "intermediate/DISSERTATION_FOR_DOCX.md"
BIB = ROOT / "paper/current_dissertation/references/REFERENCES_CURRENT.bib"
TEMPLATE_SOURCE = Path("/mnt/d/Desktop/paper/Template_MSc-Diss_v2.docx")
TEMPLATE_COPY = ROOT / "paper/templates/NTU_template.docx"
DOCX = BUILD / "DISSERTATION_NTU_TEMPLATE_DRAFT.docx"
WINDOWS_DOCX = Path("/mnt/d/Desktop/paper/DISSERTATION_NTU_TEMPLATE_DRAFT.docx")
AUDIT = BUILD / "reports/FINAL_DELIVERY_AUDIT.json"


def digest(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def citation_keys(text: str, valid_keys: set[str]) -> set[str]:
    keys = set()
    for group in re.findall(r"\[([^\]]*@[A-Za-z0-9_:.+\-/]+[^\]]*)\]", text):
        keys.update(re.findall(r"@([A-Za-z0-9_:.+\-/]+)", group))
    return keys & valid_keys


def decimal_tokens(text: str) -> Counter[str]:
    return Counter(re.findall(r"(?<![A-Za-z])[-+]?\d+\.\d+(?:[eE][-+]?\d+)?", text))


source_text = SOURCE.read_text(encoding="utf-8")
prepared_text = PREPARED.read_text(encoding="utf-8")
bib_keys = set(re.findall(r"(?m)^@[A-Za-z]+\{([^,]+),", BIB.read_text(encoding="utf-8")))
source_keys = citation_keys(source_text, bib_keys)
prepared_keys = citation_keys(prepared_text, bib_keys)

source_table_text = "\n".join(
    line
    for line in source_text.splitlines()
    if line.startswith("|") and not re.match(r"^\|(?:\s*:?-{3,}:?\s*\|)+\s*$", line)
)

document = Document(DOCX)
docx_table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
source_decimals = decimal_tokens(source_table_text)
docx_decimals = decimal_tokens(docx_table_text)

figure_paths = sorted((ROOT / "paper/current_dissertation/figures/png").glob("*.png"))
figure_hashes = {digest(path): path.name for path in figure_paths}
with ZipFile(DOCX) as archive:
    media_hashes = {
        sha256(archive.read(name)).hexdigest(): name
        for name in archive.namelist()
        if name.startswith("word/media/")
    }
    document_xml = archive.read("word/document.xml").decode("utf-8")

paragraph_texts = [paragraph.text.strip() for paragraph in document.paragraphs]
heading_one = [
    paragraph.text.strip() for paragraph in document.paragraphs if paragraph.style.name == "Heading 1"
]
reference_start = paragraph_texts.index("References")
appendix_start = paragraph_texts.index("Appendices")
reference_entries = [
    text for text in paragraph_texts[reference_start + 1 : appendix_start] if " — " in text
]

required_reports = [
    "DOCX_BUILD_REPORT.md",
    "DOCX_SELF_CHECK_REPORT.md",
    "DOCX_BUILD_STATUS.md",
]
windows_report_dir = Path("/mnt/d/Desktop/paper")
report_copy_errors = []
for name in required_reports:
    project = BUILD / name
    windows = windows_report_dir / name
    if not project.is_file() or not windows.is_file() or digest(project) != digest(windows):
        report_copy_errors.append(name)

forbidden_project_files = [
    str(path.relative_to(ROOT))
    for path in BUILD.rglob("*")
    if path.is_file() and path.suffix.lower() in {".pdf", ".tex", ".latex"}
]
forbidden_windows_files = [
    str(path)
    for path in [
        windows_report_dir / "DISSERTATION_NTU_TEMPLATE_DRAFT_QA.pdf",
        windows_report_dir / "DISSERTATION_NTU_TEMPLATE_DRAFT.pdf",
    ]
    if path.exists()
]

source_hash_expected = "90e782afa0d7d400bd45c7b80cda6fbe5a0ac67e1dfd79ccdf1b514df1dadb7c"
bib_hash_expected = "3c05b30da758ea8f60ceef4f68c9f73753818f90057d3955ec113fd6a02baf82"
template_hash_expected = "cd82dc0a40d8a8a6e3b093b144c8f3573578c83ed3f2990913a94d8ca2cdaff4"

checks = {
    "source_markdown_unchanged": digest(SOURCE) == source_hash_expected,
    "bibliography_unchanged": digest(BIB) == bib_hash_expected,
    "windows_template_unchanged": digest(TEMPLATE_SOURCE) == template_hash_expected,
    "project_template_matches_windows": digest(TEMPLATE_COPY) == digest(TEMPLATE_SOURCE),
    "project_docx_exists": DOCX.is_file(),
    "windows_docx_exists": WINDOWS_DOCX.is_file(),
    "windows_docx_matches_project": digest(DOCX) == digest(WINDOWS_DOCX),
    "docx_sha256": digest(DOCX),
    "paragraphs": len(document.paragraphs),
    "tables": len(document.tables),
    "inline_shapes": len(document.inline_shapes),
    "chapter_heading_count": sum(bool(re.match(r"^Chapter [1-6]\b", text)) for text in heading_one),
    "references_present": "References" in heading_one,
    "appendices_present": "Appendices" in heading_one,
    "reference_entries": len(reference_entries),
    "caption_paragraphs": sum(paragraph.style.name == "Caption" for paragraph in document.paragraphs),
    "source_citation_keys": len(source_keys),
    "prepared_citation_keys": len(prepared_keys),
    "citation_keys_preserved": source_keys == prepared_keys,
    "table_decimal_tokens_preserved": source_decimals == docx_decimals,
    "missing_table_decimals": dict(source_decimals - docx_decimals),
    "unexpected_table_decimals": dict(docx_decimals - source_decimals),
    "source_figure_count": len(figure_hashes),
    "source_figures_embedded_unchanged": all(hash_value in media_hashes for hash_value in figure_hashes),
    "missing_embedded_figure_hashes": [
        name for hash_value, name in figure_hashes.items() if hash_value not in media_hashes
    ],
    "toc_field_present": 'TOC \\o &quot;1-3&quot; \\h \\z \\u' in document_xml or 'TOC \\o "1-3" \\h \\z \\u' in document_xml,
    "list_of_figures_field_present": "TOC \\f F \\h \\z" in document_xml,
    "list_of_tables_field_present": "TOC \\f T \\h \\z" in document_xml,
    "figure_tc_fields": document_xml.count(" TC &quot;Figure ") + document_xml.count(' TC "Figure '),
    "table_tc_fields": document_xml.count(" TC &quot;Table ") + document_xml.count(' TC "Table '),
    "rendered_png_pages": len(list((BUILD / "rendered_pages").glob("page-*.png"))),
    "report_copy_errors": report_copy_errors,
    "forbidden_project_files": forbidden_project_files,
    "forbidden_windows_files": forbidden_windows_files,
    "status_requires_layout_fixes": "Requires DOCX layout fixes" in (BUILD / "DOCX_BUILD_STATUS.md").read_text(encoding="utf-8"),
}

failures = [
    key
    for key in [
        "source_markdown_unchanged",
        "bibliography_unchanged",
        "windows_template_unchanged",
        "project_template_matches_windows",
        "project_docx_exists",
        "windows_docx_exists",
        "windows_docx_matches_project",
        "references_present",
        "appendices_present",
        "citation_keys_preserved",
        "table_decimal_tokens_preserved",
        "source_figures_embedded_unchanged",
        "toc_field_present",
        "list_of_figures_field_present",
        "list_of_tables_field_present",
        "status_requires_layout_fixes",
    ]
    if not checks[key]
]
if checks["tables"] != 16:
    failures.append("tables")
if checks["inline_shapes"] != 11:
    failures.append("inline_shapes")
if checks["chapter_heading_count"] != 6:
    failures.append("chapter_heading_count")
if checks["reference_entries"] != 29:
    failures.append("reference_entries")
if checks["caption_paragraphs"] != 27:
    failures.append("caption_paragraphs")
if checks["figure_tc_fields"] != 11:
    failures.append("figure_tc_fields")
if checks["table_tc_fields"] != 16:
    failures.append("table_tc_fields")
if checks["rendered_png_pages"] != 84:
    failures.append("rendered_png_pages")
if report_copy_errors:
    failures.append("report_copy_errors")
if forbidden_project_files or forbidden_windows_files:
    failures.append("forbidden_outputs")

checks["audit_failures"] = failures
checks["audit_result"] = "PASS_WITH_LAYOUT_FAILURE_STATUS" if not failures else "FAIL"
AUDIT.write_text(json.dumps(checks, indent=2) + "\n", encoding="utf-8")
print(json.dumps(checks, indent=2))
if failures:
    raise SystemExit(1)
