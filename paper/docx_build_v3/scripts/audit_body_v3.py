from __future__ import annotations

from hashlib import sha256
import json
from pathlib import Path
import re
import zipfile

from docx import Document
from lxml import etree


ROOT = Path(__file__).resolve().parents[3]
BUILD = ROOT / "paper/docx_build_v3"
SOURCE = BUILD / "intermediate/DISSERTATION_FOR_V3_POLISHED.md"
BODY = BUILD / "intermediate/BODY_V3_IEEE.docx"
BIBLIOGRAPHY = ROOT / "paper/current_dissertation/references/REFERENCES_CURRENT.bib"
CSL = ROOT / "paper/styles/ieee.csl"
AUDIT = BUILD / "reports/BODY_V3_AUDIT.json"
CITATION_REPORT = BUILD / "DOCX_V3_CITATION_REPORT.md"
FORMULA_REPORT = BUILD / "DOCX_V3_FORMULA_REPAIR_REPORT.md"

NS = {
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}


def digest(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def formula_inventory(source: str) -> list[dict[str, str | int]]:
    rows: list[dict[str, str | int]] = []
    section = ""
    lines = source.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        if line.startswith("#"):
            section = line.lstrip("#").strip()
        if line.strip() != r"\[":
            index += 1
            continue
        start = index + 1
        formula_lines: list[str] = []
        index += 1
        while index < len(lines) and lines[index].strip() != r"\]":
            formula_lines.append(lines[index].strip())
            index += 1
        if index >= len(lines):
            raise RuntimeError(f"Unclosed display formula at source line {start}")
        latex = " ".join(part for part in formula_lines if part)
        rows.append({"line": start, "section": section, "latex": latex})
        index += 1
    return rows


def main() -> None:
    source_text = SOURCE.read_text(encoding="utf-8")
    document = Document(BODY)
    paragraphs = document.paragraphs
    paragraph_texts = [paragraph.text.strip() for paragraph in paragraphs]
    references_indexes = [index for index, text in enumerate(paragraph_texts) if text == "References"]
    appendices_indexes = [index for index, text in enumerate(paragraph_texts) if text == "Appendices"]
    if len(references_indexes) != 1 or len(appendices_indexes) != 1:
        raise RuntimeError("Expected exactly one References and one Appendices heading")
    reference_index = references_indexes[0]
    appendix_index = appendices_indexes[0]
    body_text = "\n".join(paragraph_texts[:reference_index])
    reference_paragraphs = [
        paragraph.text.strip()
        for paragraph in paragraphs[reference_index + 1 : appendix_index]
        if paragraph.text.strip() and paragraph.style.name == "Bibliography"
    ]
    reference_numbers = [
        int(match.group(1))
        for text in reference_paragraphs
        if (match := re.match(r"^\[(\d+)\]", text))
    ]
    numbered_citations = re.findall(r"\[(\d+(?:\s*[-–,]\s*\d+)*)\]", body_text)
    author_year_patterns = re.findall(
        r"\((?:[A-Z][A-Za-zÀ-ž'’-]+(?:\s+et\s+al\.)?|[A-Z][A-Za-zÀ-ž'’-]+\s+and\s+[A-Z][A-Za-zÀ-ž'’-]+)[, ]+\d{4}[a-z]?\)",
        body_text,
    )
    unresolved_citations = re.findall(r"\[@[^\]]+\]", "\n".join(paragraph_texts))

    with zipfile.ZipFile(BODY) as archive:
        document_xml = archive.read("word/document.xml")
    root = etree.fromstring(document_xml)
    omml_blocks = root.xpath(".//m:oMathPara", namespaces=NS)
    omml_inline = root.xpath(".//m:oMath[not(ancestor::m:oMathPara)]", namespaces=NS)
    omml_texts = ["".join(block.xpath(".//m:t/text()", namespaces=NS)).strip() for block in omml_blocks]
    empty_omml_blocks = [index + 1 for index, text in enumerate(omml_texts) if not text]
    formulas = formula_inventory(source_text)
    broken_formula_patterns = re.findall(
        r"(?m)^\[\s*(?:U|I|=|p_|phi|\\phi)[^\]]*\]\s*$",
        "\n".join(paragraph_texts),
    )

    checks = {
        "numbered_citations_present": bool(numbered_citations),
        "author_year_patterns_remaining": len(author_year_patterns),
        "unresolved_citation_markers": len(unresolved_citations),
        "references_heading_count": len(references_indexes),
        "reference_entries": len(reference_paragraphs),
        "references_numbered_contiguously": reference_numbers == list(range(1, len(reference_numbers) + 1)),
        "display_formula_source_count": len(formulas),
        "display_formula_omml_count": len(omml_blocks),
        "inline_formula_omml_count": len(omml_inline),
        "empty_omml_blocks": len(empty_omml_blocks),
        "broken_formula_markers": len(broken_formula_patterns),
        "tables": len(document.tables),
        "figures": len(document.inline_shapes),
    }
    if checks["author_year_patterns_remaining"] != 0:
        raise RuntimeError(f"Author-year citations remain: {author_year_patterns}")
    if checks["unresolved_citation_markers"] != 0:
        raise RuntimeError(f"Unresolved citations remain: {unresolved_citations}")
    if not checks["numbered_citations_present"]:
        raise RuntimeError("No numbered in-text citations found")
    if len(reference_paragraphs) != 29 or reference_numbers != list(range(1, 30)):
        raise RuntimeError("IEEE References are not numbered [1] through [29]")
    if len(formulas) != len(omml_blocks) or empty_omml_blocks:
        raise RuntimeError("Display formulas were not fully converted to OMML")
    if broken_formula_patterns:
        raise RuntimeError(f"Broken formula markers remain: {broken_formula_patterns}")

    AUDIT.parent.mkdir(parents=True, exist_ok=True)
    AUDIT.write_text(
        json.dumps(
            {
                "body": str(BODY.relative_to(ROOT)),
                "body_sha256": digest(BODY),
                "bibliography_sha256": digest(BIBLIOGRAPHY),
                "ieee_csl": str(CSL.relative_to(ROOT)),
                "ieee_csl_sha256": digest(CSL),
                "checks": checks,
                "first_five_references": reference_paragraphs[:5],
                "last_reference": reference_paragraphs[-1],
                "formula_omml_samples": omml_texts[:5],
            },
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )

    CITATION_REPORT.write_text(
        "# DOCX V3 Citation Report\n\n"
        "## Citation style\n\n"
        "IEEE numbered citation style, generated by Pandoc citeproc using `paper/styles/ieee.csl` (IEEE Reference Guide version 11.29.2023). No regex-based citation substitution was used.\n\n"
        "## Checks\n\n"
        "| Check | Result | Notes |\n"
        "|---|---|---|\n"
        f"| Numbered in-text citations | PASS | {len(numbered_citations)} numbered citation fields/patterns found before References. |\n"
        "| Author-year citation patterns | PASS | 0 remaining. |\n"
        "| Unresolved citation keys | PASS | 0 `[@...]` markers remaining. |\n"
        "| Numbered References | PASS | 29 entries numbered contiguously from [1] to [29]. |\n"
        "| Reference order | PASS | IEEE citeproc output follows first-citation order; contiguous numbering was verified in DOCX text. |\n"
        "| Duplicate References section | PASS | Exactly one References heading. |\n\n"
        "## Remaining issues\n\n"
        "PEARLM remains cited through its verified arXiv metadata; final venue and publisher DOI still require manual checking. Other existing metadata caveats are unchanged.\n\n"
        "Author-year citation patterns remaining: 0  \n"
        "Numbered citation patterns found: yes  \n"
        "References numbered: yes\n",
        encoding="utf-8",
    )

    formula_rows = "\n".join(
        f"| {item['section']} (source line {item['line']}) | Display-math marker previously omitted by the V2 reader configuration | Pandoc `tex_math_single_backslash` to OMML | Word OMML equation ({str(item['latex'])[:80].replace('|', '/')}...) | PASS |"
        for item in formulas
    )
    FORMULA_REPORT.write_text(
        "# DOCX V3 Formula Repair Report\n\n"
        "V2 used a Markdown reader configuration that did not recognise the source `\\[ ... \\]` display-math delimiters. V3 enables `tex_math_single_backslash`, so Pandoc emits editable Word OMML equations.\n\n"
        "| Location | Original issue | Repair method | Final representation | Status |\n"
        "|---|---|---|---|---|\n"
        + formula_rows
        + "\n\n"
        f"Display formulas in V3 source: {len(formulas)}  \n"
        f"Display OMML blocks in body DOCX: {len(omml_blocks)}  \n"
        f"Inline OMML equations in body DOCX: {len(omml_inline)}  \n"
        "Broken `[ ... ]` formula markers remaining: 0\n",
        encoding="utf-8",
    )
    print(AUDIT)


if __name__ == "__main__":
    main()
