from __future__ import annotations

from hashlib import sha256
import json
from pathlib import Path
import re
import sys
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[3]
SOURCE = ROOT / "paper/current_dissertation/FULL_DISSERTATION_CURRENT.md"
BIBLIOGRAPHY = ROOT / "paper/current_dissertation/references/REFERENCES_CURRENT.bib"
TEMPLATE = ROOT / "paper/templates/NTU_template.docx"
DOCX = ROOT / "paper/docx_build_v3/DISSERTATION_NTU_TEMPLATE_DRAFT_V3.docx"
LOGO = ROOT / "paper/docx_build_v3/intermediate/template_logo.jpg"
ASSEMBLY_AUDIT = ROOT / "paper/docx_build_v3/reports/ASSEMBLY_AUDIT.json"
BODY_AUDIT = ROOT / "paper/docx_build_v3/reports/BODY_V3_AUDIT.json"
SOURCE_AUDIT = ROOT / "paper/docx_build_v3/reports/SOURCE_TRANSFORMATION_AUDIT.json"
ITERATION = sys.argv[1] if len(sys.argv) > 1 else "iteration4"
RENDER_AUDIT = ROOT / "paper/docx_build_v3/reports" / f"RENDER_AUDIT_{ITERATION.upper()}.json"
RENDER_TEXT = Path(sys.argv[2]) if len(sys.argv) > 2 else Path(
    f"/tmp/v3_render_{ITERATION}/DISSERTATION_NTU_TEMPLATE_DRAFT_V3.txt"
)
PAGINATION_MAP = ROOT / "paper/docx_build_v3/intermediate/PAGINATION_MAP.json"
OUTPUT = ROOT / "paper/docx_build_v3/reports/FINAL_AUDIT.json"

EXPECTED_HASHES = {
    "source": "90e782afa0d7d400bd45c7b80cda6fbe5a0ac67e1dfd79ccdf1b514df1dadb7c",
    "bibliography": "3c05b30da758ea8f60ceef4f68c9f73753818f90057d3955ec113fd6a02baf82",
    "template": "cd82dc0a40d8a8a6e3b093b144c8f3573578c83ed3f2990913a94d8ca2cdaff4",
}


def digest(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def source_citation_keys(text: str) -> set[str]:
    keys: set[str] = set()
    for group in re.findall(r"\[([^\]]*@[a-z][A-Za-z0-9_:.+\-/]+[^\]]*)\]", text):
        keys.update(re.findall(r"@([a-z][A-Za-z0-9_:.+\-/]+)", group))
    return keys


def reference_numbers(reference_entries: list[str]) -> list[int]:
    numbers: list[int] = []
    for entry in reference_entries:
        match = re.match(r"^\[(\d+)\]", entry)
        if match:
            numbers.append(int(match.group(1)))
    return numbers


def table_xml_checks(document: Document) -> dict[str, object]:
    widths: list[int] = []
    repeat_header_count = 0
    vertical_borders_removed = 0
    for table in document.tables:
        properties = table._tbl.tblPr
        width = properties.find(qn("w:tblW"))
        if width is not None and width.get(qn("w:type")) == "dxa":
            widths.append(int(width.get(qn("w:w"))))
        borders = properties.find(qn("w:tblBorders"))
        if borders is not None:
            removed = True
            for edge in ("left", "right", "insideV"):
                border = borders.find(qn(f"w:{edge}"))
                removed = removed and border is not None and border.get(qn("w:val")) == "nil"
            if removed:
                vertical_borders_removed += 1
        if table.rows:
            header = table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader"))
            if header is not None and header.get(qn("w:val")) == "true":
                repeat_header_count += 1
    return {
        "table_widths": widths,
        "table_widths_within_textblock": bool(widths) and all(width <= 8640 for width in widths),
        "repeat_header_rows": repeat_header_count,
        "vertical_borders_removed_tables": vertical_borders_removed,
    }


def element_text(element) -> str:
    if element.tag != qn("w:p"):
        return ""
    return "".join(node.text or "" for node in element.xpath(".//w:t")).strip()


def table_caption_position_checks(document: Document) -> dict[str, object]:
    below = 0
    above_neighbor = 0
    above_same_table = 0
    for table in document.tables:
        below_id = None
        next_node = table._tbl.getnext()
        while next_node is not None and next_node.tag == qn("w:p"):
            text = element_text(next_node)
            if text:
                match = re.match(r"^Table\s+(\d+\.\d+[a-z]?)\.", text)
                if match:
                    below += 1
                    below_id = match.group(1)
                break
            next_node = next_node.getnext()

        prev_node = table._tbl.getprevious()
        while prev_node is not None and prev_node.tag == qn("w:p"):
            text = element_text(prev_node)
            if text:
                match = re.match(r"^Table\s+(\d+\.\d+[a-z]?)\.", text)
                if match:
                    above_neighbor += 1
                    if below_id is not None and match.group(1) == below_id:
                        above_same_table += 1
                break
            prev_node = prev_node.getprevious()
    return {
        "table_captions_below_tables": below,
        "table_caption_neighbors_above_tables": above_neighbor,
        "table_captions_above_same_table": above_same_table,
        "all_table_captions_below": below == len(document.tables) and above_same_table == 0,
    }


source_text = SOURCE.read_text(encoding="utf-8")
source_keys = source_citation_keys(source_text)
document = Document(DOCX)
assembly_audit = json.loads(ASSEMBLY_AUDIT.read_text(encoding="utf-8"))
body_audit = json.loads(BODY_AUDIT.read_text(encoding="utf-8"))
source_audit = json.loads(SOURCE_AUDIT.read_text(encoding="utf-8"))
render_audit = json.loads(RENDER_AUDIT.read_text(encoding="utf-8"))
pagination_map = json.loads(PAGINATION_MAP.read_text(encoding="utf-8"))

paragraph_texts = [paragraph.text.strip() for paragraph in document.paragraphs]
all_text = "\n".join(paragraph_texts)
heading_one = [
    paragraph.text.strip()
    for paragraph in document.paragraphs
    if paragraph.style.name == "Heading 1"
]
reference_start = next(
    index
    for index, paragraph in enumerate(document.paragraphs)
    if paragraph.text.strip() == "References" and paragraph.style.name == "Heading 1"
)
appendix_start = next(
    index
    for index, paragraph in enumerate(document.paragraphs)
    if paragraph.text.strip() == "Appendices" and paragraph.style.name == "Heading 1"
)
reference_entries = [
    paragraph.text.strip()
    for paragraph in document.paragraphs[reference_start + 1 : appendix_start]
    if paragraph.text.strip() and paragraph.style.name == "Reference"
]

contents_paragraphs = [
    paragraph
    for paragraph in document.paragraphs
    if paragraph.style.name in {"Contents 1", "Contents 2"}
]
list_paragraphs = [
    paragraph for paragraph in document.paragraphs if paragraph.style.name == "List Entry"
]

section_page_number_formats = []
for section in document.sections:
    page_number_type = section._sectPr.find(qn("w:pgNumType"))
    if page_number_type is not None:
        section_page_number_formats.append(
            {
                "format": page_number_type.get(qn("w:fmt")),
                "start": page_number_type.get(qn("w:start")),
            }
        )

with ZipFile(DOCX) as archive:
    zip_integrity_error = archive.testzip()
    media = {
        sha256(archive.read(name)).hexdigest(): name
        for name in archive.namelist()
        if name.startswith("word/media/")
    }
    document_xml = archive.read("word/document.xml").decode("utf-8")
    styles_xml = archive.read("word/styles.xml").decode("utf-8")
    footer_xml = "\n".join(
        archive.read(name).decode("utf-8")
        for name in archive.namelist()
        if re.match(r"word/footer\d+\.xml$", name)
    )

source_figure_hashes = {
    image["sha256"]: image["prepared"]
    for image in source_audit.get("images", [])
}
missing_source_figures = [
    path for hash_value, path in source_figure_hashes.items() if hash_value not in media
]
table_checks = table_xml_checks(document)
caption_position_checks = table_caption_position_checks(document)
reference_nums = reference_numbers(reference_entries)
render_long_decimal_pages = render_audit["long_decimal_pages"]
references_first_page = render_audit["references_pages"][0]
appendices_first_page = render_audit["appendices_pages"][0]
rendered_text = RENDER_TEXT.read_text(encoding="utf-8", errors="replace") if RENDER_TEXT.is_file() else ""
render_repository_paths = sorted(
    set(re.findall(r"(?:paper|reports|docs|scripts|thesis_analysis_pack)/[^\s,;.)]+", rendered_text))
)
short_figure_alt_phrases = [
    "Figure 3.1 framework overview",
    "Figure 3.2 alpha-sweep design",
    "Figure 3.3 validation gate",
    "Figure 4.1 strict NDCG comparison",
    "Figure 4.2 explanation metric endpoints",
    "Figure 4.3 LIR-oriented trade-off",
    "Figure 4.4 SEP-oriented trade-off",
    "Figure 4.5 ETD-oriented trade-off",
    "Figure 5.1 PGPR/UCPR ablation",
    "Figure 5.2 experiment-status matrix",
    "Figure C.1 Amazon-Book KGAT boundary decision flow",
]
short_figure_alt_remaining = [
    phrase for phrase in short_figure_alt_phrases if phrase in rendered_text
]

checks = {
    "source_markdown_unchanged": digest(SOURCE) == EXPECTED_HASHES["source"],
    "bibliography_unchanged": digest(BIBLIOGRAPHY) == EXPECTED_HASHES["bibliography"],
    "template_unchanged": digest(TEMPLATE) == EXPECTED_HASHES["template"],
    "docx_exists": DOCX.is_file(),
    "docx_sha256": digest(DOCX),
    "zip_integrity": zip_integrity_error is None,
    "template_used_as_mother_document": assembly_audit["template_used_as_mother_document"],
    "pagination_map_applied": assembly_audit["pagination_map_applied"],
    "paragraphs": len(document.paragraphs),
    "tables": len(document.tables),
    "inline_shapes_total": len(document.inline_shapes),
    "body_figure_count": len(document.inline_shapes) - 1,
    "sections": len(document.sections),
    "landscape_sections": sum(
        section.orientation == WD_ORIENT.LANDSCAPE for section in document.sections
    ),
    "chapter_heading_count": sum(
        bool(re.match(r"^Chapter [1-6]\b", text)) for text in heading_one
    ),
    "references_present": "References" in heading_one,
    "appendices_present": "Appendices" in heading_one,
    "reference_entries": len(reference_entries),
    "references_numbered_contiguously": reference_nums == list(range(1, 30)),
    "source_citation_key_count": len(source_keys),
    "unresolved_citation_syntax_absent": not bool(re.search(r"\[@[a-z]", all_text)),
    "author_year_patterns_absent": not bool(
        re.search(
            r"\((?:Guo et al\.\s+2022|Zhang and Chen\s+2020|Balloccu et al\.\s+2022[a-z]?)\)",
            all_text,
        )
    ),
    "numbered_citations_present": bool(re.search(r"\[[0-9]+(?:\s*[-,]\s*[0-9]+)*\]", all_text)),
    "figure_caption_count": sum(
        paragraph.style.name == "Caption" and paragraph.text.strip().startswith("Figure ")
        for paragraph in document.paragraphs
    ),
    "table_caption_count": sum(
        paragraph.style.name == "Caption" and paragraph.text.strip().startswith("Table ")
        for paragraph in document.paragraphs
    ),
    "source_figures_embedded_unchanged": not missing_source_figures,
    "missing_source_figures": missing_source_figures,
    "template_logo_embedded_unchanged": digest(LOGO) in media,
    "media_part_count": len(media),
    "media_part_count_at_least_required": len(media) >= 12,
    "toc_entries": len(contents_paragraphs),
    "list_entries": len(list_paragraphs),
    "toc_page_numbers_complete": all(
        paragraph.text.split("\t")[-1].strip() for paragraph in contents_paragraphs
    ),
    "list_page_numbers_complete": all(
        paragraph.text.split("\t")[-1].strip() for paragraph in list_paragraphs
    ),
    "pagination_map_heading_count": pagination_map["mapped_heading_count"],
    "pagination_map_figure_count": pagination_map["mapped_figure_count"],
    "pagination_map_table_count": pagination_map["mapped_table_count"],
    "page_number_formats": section_page_number_formats,
    "page_field_count": footer_xml.count(" PAGE "),
    "figure_tc_fields": document_xml.count(' TC "Figure ') + document_xml.count(" TC &quot;Figure "),
    "table_tc_fields": document_xml.count(' TC "Table ') + document_xml.count(" TC &quot;Table "),
    "caption_style_explicit_black": 'w:styleId="Caption"' in styles_xml and 'w:color w:val="000000"' in styles_xml,
    "display_formula_source_count": body_audit["checks"]["display_formula_source_count"],
    "display_formula_omml_count": body_audit["checks"]["display_formula_omml_count"],
    "inline_formula_omml_count": body_audit["checks"]["inline_formula_omml_count"],
    "empty_omml_blocks": body_audit["checks"]["empty_omml_blocks"],
    "broken_formula_markers": body_audit["checks"]["broken_formula_markers"],
    "numeric_long_decimals_remaining_in_intermediate": source_audit["long_decimals_remaining"],
    "numeric_changes": source_audit["numeric_changes"],
    "citation_keys_preserved": source_audit["citation_keys_preserved"],
    "source_markdown_modified_by_transform": source_audit["source_markdown_modified"],
    "bibliography_modified_by_transform": source_audit["bibliography_modified"],
    "render_page_count": render_audit["page_count"],
    "render_portrait_pages": render_audit["portrait_pages"],
    "render_landscape_pages": render_audit["landscape_pages"],
    "render_blank_pages": render_audit["blank_pages"],
    "render_edge_touch_pages": render_audit["edge_touch_pages"],
    "render_replacement_character_pages": render_audit["replacement_character_pages"],
    "render_image_error_marker_pages": render_audit["image_error_marker_pages"],
    "render_author_year_pages": render_audit["author_year_pages"],
    "render_broken_formula_pages": render_audit["broken_formula_pages"],
    "render_table_margin_risk_pages": render_audit["table_margin_risk_pages"],
    "render_long_decimal_pages": render_long_decimal_pages,
    "render_long_decimal_pages_references_only": all(
        references_first_page <= page < appendices_first_page for page in render_long_decimal_pages
    ),
    "render_contact_sheets": render_audit["contact_sheets"],
    "render_repository_paths": render_repository_paths,
    "render_repository_paths_absent": not render_repository_paths,
    "short_figure_alt_remaining": short_figure_alt_remaining,
    "single_visible_figure_caption_set": not short_figure_alt_remaining,
}
checks.update(table_checks)
checks.update(caption_position_checks)

required_boolean_checks = [
    "source_markdown_unchanged",
    "bibliography_unchanged",
    "template_unchanged",
    "docx_exists",
    "zip_integrity",
    "template_used_as_mother_document",
    "pagination_map_applied",
    "references_present",
    "appendices_present",
    "references_numbered_contiguously",
    "unresolved_citation_syntax_absent",
    "author_year_patterns_absent",
    "numbered_citations_present",
    "source_figures_embedded_unchanged",
    "template_logo_embedded_unchanged",
    "media_part_count_at_least_required",
    "toc_page_numbers_complete",
    "list_page_numbers_complete",
    "caption_style_explicit_black",
    "citation_keys_preserved",
    "render_long_decimal_pages_references_only",
    "table_widths_within_textblock",
    "all_table_captions_below",
    "render_repository_paths_absent",
    "single_visible_figure_caption_set",
]
failures = [key for key in required_boolean_checks if not checks[key]]

expected_counts = {
    "tables": 17,
    "inline_shapes_total": 12,
    "body_figure_count": 11,
    "sections": 3,
    "landscape_sections": 0,
    "chapter_heading_count": 6,
    "reference_entries": 29,
    "source_citation_key_count": 29,
    "figure_caption_count": 11,
    "table_caption_count": 17,
    "toc_entries": 49,
    "list_entries": 28,
    "pagination_map_heading_count": 49,
    "pagination_map_figure_count": 11,
    "pagination_map_table_count": 17,
    "page_field_count": 2,
    "figure_tc_fields": 11,
    "table_tc_fields": 17,
    "display_formula_source_count": 31,
    "display_formula_omml_count": 31,
    "empty_omml_blocks": 0,
    "broken_formula_markers": 0,
    "numeric_long_decimals_remaining_in_intermediate": 0,
    "render_page_count": 73,
    "render_portrait_pages": 73,
    "render_landscape_pages": 0,
    "repeat_header_rows": 17,
    "vertical_borders_removed_tables": 17,
    "table_captions_below_tables": 17,
    "table_captions_above_same_table": 0,
}
for key, expected in expected_counts.items():
    if checks[key] != expected:
        failures.append(key)

for key in [
    "render_blank_pages",
    "render_edge_touch_pages",
    "render_replacement_character_pages",
    "render_image_error_marker_pages",
    "render_author_year_pages",
    "render_broken_formula_pages",
    "render_table_margin_risk_pages",
    "missing_source_figures",
    "render_repository_paths",
    "short_figure_alt_remaining",
]:
    if checks[key]:
        failures.append(key)

if checks["page_number_formats"] != [
    {"format": "lowerRoman", "start": "1"},
    {"format": "decimal", "start": "1"},
]:
    failures.append("page_number_formats")

for key, value in body_audit["checks"].items():
    if isinstance(value, bool) and not value:
        failures.append(f"body_audit.{key}")

checks["audit_failures"] = sorted(set(failures))
checks["audit_result"] = "PASS" if not checks["audit_failures"] else "FAIL"
OUTPUT.write_text(json.dumps(checks, indent=2, default=dict) + "\n", encoding="utf-8")
print(json.dumps(checks, indent=2, default=dict))
if checks["audit_failures"]:
    raise SystemExit(1)
