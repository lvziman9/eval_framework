from __future__ import annotations

from copy import deepcopy
from hashlib import sha256
from io import BytesIO
import json
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[3]
TEMPLATE = ROOT / "paper/templates/NTU_template.docx"
RAW_BODY = ROOT / "paper/docx_build_v2/intermediate/DISSERTATION_BODY_PANDOC_RAW.docx"
LOGO = ROOT / "paper/docx_build_v2/intermediate/template_media/image1.jpg"
PAGINATION_MAP = ROOT / "paper/docx_build_v2/intermediate/PAGINATION_MAP.json"
OUTPUT = ROOT / "paper/docx_build_v2/DISSERTATION_NTU_TEMPLATE_DRAFT_V2.docx"
AUDIT = ROOT / "paper/docx_build_v2/reports/ASSEMBLY_AUDIT.json"

TITLE = "A Canonical Framework for Accuracy–Explainability Trade-off Analysis in Native-Path Knowledge Graph Recommenders"
AUTHOR = "[Name to be filled]"
SUPERVISOR = "[Supervisor to be filled]"
SCHOOL = "School of Computer Science and Engineering"
DEGREE = "Master of Science"
DATE = "[Submission date to be filled]"
YEAR = "[Submission year to be filled]"


def digest(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def set_font_properties(target, name: str, size: float, bold=None, italic=None) -> None:
    target.font.name = name
    target.font.size = Pt(size)
    if bold is not None:
        target.font.bold = bold
    if italic is not None:
        target.font.italic = italic
    run_properties = target.element.get_or_add_rPr()
    fonts = run_properties.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), name)


def get_or_add_style(document, name: str, style_type: WD_STYLE_TYPE, base=None):
    try:
        style = document.styles[name]
    except KeyError:
        style = document.styles.add_style(name, style_type)
    if base is not None:
        style.base_style = document.styles[base]
    return style


def configure_styles(document) -> None:
    normal = document.styles["Normal"]
    set_font_properties(normal, "Times New Roman", 12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True

    title = document.styles["Title"]
    set_font_properties(title, "Times New Roman", 16.5, bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)

    subtitle = document.styles["Subtitle"]
    set_font_properties(subtitle, "Times New Roman", 12, bold=True)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)

    heading_1 = document.styles["Heading 1"]
    set_font_properties(heading_1, "Times New Roman", 20, bold=True)
    heading_1.font.color.rgb = RGBColor(0, 0, 0)
    heading_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_1.paragraph_format.space_before = Pt(0)
    heading_1.paragraph_format.space_after = Pt(14)
    heading_1.paragraph_format.keep_with_next = True
    heading_1.paragraph_format.page_break_before = True

    heading_2 = document.styles["Heading 2"]
    set_font_properties(heading_2, "Times New Roman", 14, bold=True)
    heading_2.font.color.rgb = RGBColor(0, 0, 0)
    heading_2.paragraph_format.space_before = Pt(14)
    heading_2.paragraph_format.space_after = Pt(6)
    heading_2.paragraph_format.keep_with_next = True
    heading_2.paragraph_format.page_break_before = False

    heading_3 = document.styles["Heading 3"]
    set_font_properties(heading_3, "Times New Roman", 12, bold=True)
    heading_3.font.color.rgb = RGBColor(0, 0, 0)
    heading_3.paragraph_format.space_before = Pt(10)
    heading_3.paragraph_format.space_after = Pt(4)
    heading_3.paragraph_format.keep_with_next = True

    list_style = document.styles["List Paragraph"]
    set_font_properties(list_style, "Times New Roman", 12)
    list_style.paragraph_format.line_spacing = 1.5
    list_style.paragraph_format.space_after = Pt(3)

    front_heading = get_or_add_style(document, "Front Matter Heading", WD_STYLE_TYPE.PARAGRAPH, "Heading 1")
    set_font_properties(front_heading, "Times New Roman", 20, bold=True)
    front_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front_heading.paragraph_format.page_break_before = False
    front_heading.paragraph_format.space_after = Pt(18)

    form_heading = get_or_add_style(document, "Form Heading", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    set_font_properties(form_heading, "Times New Roman", 13.5, bold=True)
    form_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    form_heading.paragraph_format.space_after = Pt(18)

    caption = get_or_add_style(document, "Caption", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    set_font_properties(caption, "Times New Roman", 10.5)
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_together = True
    caption.paragraph_format.widow_control = True

    reference = get_or_add_style(document, "Reference", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    set_font_properties(reference, "Times New Roman", 10)
    reference.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    reference.paragraph_format.left_indent = Inches(0.3)
    reference.paragraph_format.first_line_indent = Inches(-0.3)
    reference.paragraph_format.line_spacing = 1.0
    reference.paragraph_format.space_after = Pt(4)

    toc_1 = get_or_add_style(document, "Contents 1", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    set_font_properties(toc_1, "Times New Roman", 10.5)
    toc_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_1.paragraph_format.line_spacing = 1.0
    toc_1.paragraph_format.space_after = Pt(3)
    toc_1.paragraph_format.left_indent = Inches(0)

    toc_2 = get_or_add_style(document, "Contents 2", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    set_font_properties(toc_2, "Times New Roman", 10)
    toc_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_2.paragraph_format.line_spacing = 1.0
    toc_2.paragraph_format.space_after = Pt(2)
    toc_2.paragraph_format.left_indent = Inches(0.3)

    list_entry = get_or_add_style(document, "List Entry", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    set_font_properties(list_entry, "Times New Roman", 10)
    list_entry.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    list_entry.paragraph_format.line_spacing = 1.0
    list_entry.paragraph_format.space_after = Pt(3)


def set_section_geometry(section, landscape: bool = False) -> None:
    section.page_width = Inches(11 if landscape else 8.5)
    section.page_height = Inches(8.5 if landscape else 11)
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
    else:
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    section.gutter = Inches(0)


def clear_body(document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, result, end])
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def set_page_numbering(section, fmt: str, start: int | None = None) -> None:
    section_properties = section._sectPr
    page_number_type = section_properties.find(qn("w:pgNumType"))
    if page_number_type is None:
        page_number_type = OxmlElement("w:pgNumType")
        insertion_anchor = section_properties.find(qn("w:cols"))
        if insertion_anchor is None:
            insertion_anchor = section_properties.find(qn("w:docGrid"))
        if insertion_anchor is None:
            section_properties.append(page_number_type)
        else:
            insertion_anchor.addprevious(page_number_type)
    page_number_type.set(qn("w:fmt"), fmt)
    if start is None:
        page_number_type.attrib.pop(qn("w:start"), None)
    else:
        page_number_type.set(qn("w:start"), str(start))


def configure_numbered_footer(section, fmt: str, start: int) -> None:
    section.footer.is_linked_to_previous = False
    footer_paragraph = section.footer.paragraphs[0]
    clear_paragraph(footer_paragraph)
    add_page_field(footer_paragraph)
    set_page_numbering(section, fmt, start)


def continue_page_numbering(section) -> None:
    page_number_type = section._sectPr.find(qn("w:pgNumType"))
    if page_number_type is not None:
        section._sectPr.remove(page_number_type)


def add_centered_text(document, text: str, size: float, bold=False, before=0, after=0):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    return paragraph


def add_signature_lines(document, left_label: str, right_label: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(42)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run("______________________________\t\t______________________________")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(0)
    tabs = paragraph.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(3.2))
    run = paragraph.add_run(f"{left_label}\t{right_label}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def add_form_page(document, title: str, paragraphs: list[str], signature_labels: tuple[str, str]) -> None:
    document.add_page_break()
    heading = document.add_paragraph(title, style="Form Heading")
    heading.paragraph_format.space_before = Pt(6)
    for text in paragraphs:
        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(8)
    add_signature_lines(document, *signature_labels)


def add_cover_and_forms(document, template_texts: list[str]) -> None:
    logo_paragraph = document.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_paragraph.paragraph_format.space_after = Pt(52)
    logo_paragraph.add_run().add_picture(str(LOGO), width=Inches(3.55))
    add_centered_text(document, TITLE, 16.5, bold=True, after=122)
    add_centered_text(document, AUTHOR, 13, bold=True, after=8)
    add_centered_text(document, SCHOOL.upper(), 12, bold=True, after=8)
    add_centered_text(document, YEAR, 12, bold=True)

    document.add_page_break()
    add_centered_text(document, TITLE, 16.5, bold=True, before=18, after=90)
    add_centered_text(document, AUTHOR, 13, bold=True, after=12)
    add_centered_text(document, f"Supervisor: {SUPERVISOR}", 12, bold=False, after=30)
    add_centered_text(document, SCHOOL.upper(), 12, bold=True, after=22)
    add_centered_text(
        document,
        "A DISSERTATION SUBMITTED IN PARTIAL FULFILMENT OF\nTHE REQUIREMENTS FOR THE DEGREE OF",
        12,
        bold=True,
        after=14,
    )
    add_centered_text(document, DEGREE.upper(), 12, bold=True, after=52)
    add_centered_text(document, DATE, 12, bold=True)

    originality_paragraphs = [template_texts[index] for index in [25, 27, 28, 29, 30, 31]]
    add_form_page(
        document,
        "Statement of Originality",
        originality_paragraphs,
        ("Date", "Candidate name and signature"),
    )
    add_form_page(
        document,
        "Supervisor Declaration Statement",
        [template_texts[39]],
        ("Date", "Supervisor name and signature"),
    )
    authorship_paragraphs = [template_texts[index] for index in [46, 47, 48, 49]]
    authorship_paragraphs.append(
        "[Complete this statement using the approved NTU wording. Remove the unselected option and add only verified publication and contribution details.]"
    )
    add_form_page(
        document,
        "Authorship Attribution Statement",
        authorship_paragraphs,
        ("Date", "Candidate name and signature"),
    )


def set_right_tab_with_dots(paragraph, position_inches: float) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    tabs = paragraph_properties.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        paragraph_properties.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), str(int(position_inches * 1440)))
    tabs.append(tab)


def add_static_index(document, heading: str, entries: list[tuple[str, int]], page_map: dict[str, str]) -> None:
    paragraph = document.add_paragraph(heading, style="Front Matter Heading")
    paragraph.paragraph_format.space_before = Pt(0)
    for text, level in entries:
        style = "Contents 1" if level == 1 else "Contents 2"
        entry = document.add_paragraph(style=style)
        set_right_tab_with_dots(entry, 5.85)
        entry.add_run(text)
        entry.add_run("\t")
        entry.add_run(page_map.get(text, ""))


def add_list_index(document, heading: str, entries: list[str], page_map: dict[str, str]) -> None:
    document.add_paragraph(heading, style="Front Matter Heading")
    for text in entries:
        entry = document.add_paragraph(style="List Entry")
        set_right_tab_with_dots(entry, 5.85)
        entry.add_run(text)
        entry.add_run("\t")
        identifier = re.match(r"^(Figure|Table)\s+((?:\d+\.)+\d+|[A-Z]\.\d+)", text)
        key = identifier.group(0) if identifier else text
        entry.add_run(page_map.get(key, ""))


def direct_child_text(element) -> str:
    if element.tag != qn("w:p"):
        return ""
    return "".join(node.text or "" for node in element.xpath(".//w:t")).strip()


def paragraph_style_id(element) -> str:
    if element.tag != qn("w:p"):
        return ""
    styles = element.xpath("./w:pPr/w:pStyle")
    return styles[0].get(qn("w:val")) if styles else "Normal"


def copy_numbering(source_document, destination_document) -> None:
    destination_document.part.numbering_part._element = deepcopy(
        source_document.part.numbering_part._element
    )


def remap_relationships(element, source_document, destination_document) -> None:
    mapping: dict[str, str] = {}
    relationship_attributes = [qn("r:id"), qn("r:embed"), qn("r:link")]
    for node in element.iter():
        for attribute in relationship_attributes:
            old_relationship_id = node.get(attribute)
            if not old_relationship_id:
                continue
            if old_relationship_id not in mapping:
                relationship = source_document.part.rels[old_relationship_id]
                if relationship.reltype == RT.IMAGE:
                    new_relationship_id, _ = destination_document.part.get_or_add_image(
                        BytesIO(relationship.target_part.blob)
                    )
                elif relationship.reltype == RT.HYPERLINK and relationship.is_external:
                    new_relationship_id = destination_document.part.relate_to(
                        relationship.target_ref,
                        RT.HYPERLINK,
                        is_external=True,
                    )
                else:
                    raise RuntimeError(
                        f"Unsupported body relationship: {relationship.reltype} ({old_relationship_id})"
                    )
                mapping[old_relationship_id] = new_relationship_id
            node.set(attribute, mapping[old_relationship_id])


def append_imported_element(element, source_document, destination_document):
    imported = deepcopy(element)
    remap_relationships(imported, source_document, destination_document)
    destination_document._element.body.insert(-1, imported)
    if imported.tag == qn("w:p"):
        return Paragraph(imported, destination_document._body)
    if imported.tag == qn("w:tbl"):
        return Table(imported, destination_document._body)
    return imported


def append_rebuilt_table(source_table: Table, destination_document) -> Table:
    row_count = len(source_table.rows)
    column_count = len(source_table.columns)
    rebuilt = destination_document.add_table(rows=row_count, cols=column_count)
    for row_index, source_row in enumerate(source_table.rows):
        for column_index, source_cell in enumerate(source_row.cells):
            rebuilt.cell(row_index, column_index).text = source_cell.text
    return rebuilt


def reorder_table_captions(elements: list) -> list:
    reordered = []
    index = 0
    while index < len(elements):
        element = elements[index]
        if element.tag == qn("w:tbl") and index + 1 < len(elements):
            next_element = elements[index + 1]
            next_text = direct_child_text(next_element)
            if re.match(r"^Table\s+(?:\d+(?:\.\d+)+|[A-Z]\.\d+)\.\s+", next_text):
                reordered.extend([next_element, element])
                index += 2
                continue
        reordered.append(element)
        index += 1
    return reordered


def append_tc_field(paragraph, caption_text: str, identifier: str) -> None:
    safe_text = caption_text.replace('"', "'")
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)
    paragraph._p.append(begin_run)

    instruction_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    vanish = OxmlElement("w:vanish")
    run_properties.append(vanish)
    instruction_run.append(run_properties)
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f' TC "{safe_text}" \\f {identifier} \\l 1 '
    instruction_run.append(instruction)
    paragraph._p.append(instruction_run)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph._p.append(end_run)


def set_repeat_table_header(row) -> None:
    table_row_properties = row._tr.get_or_add_trPr()
    header = table_row_properties.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        table_row_properties.append(header)
    header.set(qn("w:val"), "true")
    no_split = table_row_properties.find(qn("w:cantSplit"))
    if no_split is None:
        table_row_properties.append(OxmlElement("w:cantSplit"))


def set_cell_shading(cell, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    table_properties = table._tbl.tblPr
    borders = table_properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4" if edge.startswith("inside") else "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "666666" if edge.startswith("inside") else "000000")


def set_cell_margins(cell, twips: int = 65) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)
    for side in ("top", "left", "bottom", "right"):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(twips))
        element.set(qn("w:type"), "dxa")


def allocate_widths(table, total_width_twips: int) -> list[int]:
    scores: list[float] = []
    for column in table.columns:
        texts = [cell.text.strip() for cell in column.cells]
        longest_word = max((len(word) for text in texts for word in re.findall(r"\S+", text)), default=1)
        average_length = sum(len(text) for text in texts) / max(len(texts), 1)
        scores.append(max(8.0, min(longest_word, 38) + min(average_length ** 0.5, 14)))
    minimum = max(620, min(1200, total_width_twips // max(len(scores) * 2, 1)))
    minimum_total = minimum * len(scores)
    if minimum_total > total_width_twips:
        minimum = max(420, total_width_twips // len(scores) - 10)
        minimum_total = minimum * len(scores)
    remaining = max(0, total_width_twips - minimum_total)
    score_total = sum(scores)
    widths = [minimum + int(remaining * score / score_total) for score in scores]
    widths[-1] += total_width_twips - sum(widths)
    return widths


def apply_table_widths(table, widths: list[int]) -> None:
    table_properties = table._tbl.tblPr
    layout = table_properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
    if len(grid_columns) != len(widths):
        raise RuntimeError("Table grid does not match the logical column count")
    for grid_column, width in zip(grid_columns, widths):
        grid_column.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def table_requires_landscape(table) -> bool:
    column_count = len(table.columns)
    texts = [cell.text.strip() for row in table.rows for cell in row.cells]
    maximum_cell_length = max(map(len, texts), default=0)
    maximum_word_length = max(
        (len(word) for text in texts for word in re.findall(r"\S+", text)),
        default=0,
    )
    return (
        column_count >= 4
        or (column_count == 3 and maximum_cell_length > 180)
        or (column_count == 3 and maximum_word_length > 34)
    )


def section_properties(landscape: bool) -> OxmlElement:
    properties = OxmlElement("w:sectPr")
    section_type = OxmlElement("w:type")
    section_type.set(qn("w:val"), "nextPage")
    properties.append(section_type)
    page_size = OxmlElement("w:pgSz")
    page_size.set(qn("w:w"), "15840" if landscape else "12240")
    page_size.set(qn("w:h"), "12240" if landscape else "15840")
    if landscape:
        page_size.set(qn("w:orient"), "landscape")
    properties.append(page_size)
    page_margins = OxmlElement("w:pgMar")
    if landscape:
        margins = {"top": 1080, "right": 1080, "bottom": 1080, "left": 1080}
    else:
        margins = {"top": 1440, "right": 1440, "bottom": 1440, "left": 2160}
    for name, value in margins.items():
        page_margins.set(qn(f"w:{name}"), str(value))
    page_margins.set(qn("w:header"), "720")
    page_margins.set(qn("w:footer"), "720")
    page_margins.set(qn("w:gutter"), "0")
    properties.append(page_margins)
    columns = OxmlElement("w:cols")
    columns.set(qn("w:space"), "720")
    properties.append(columns)
    grid = OxmlElement("w:docGrid")
    grid.set(qn("w:linePitch"), "360")
    properties.append(grid)
    return properties


def make_section_break(landscape: bool) -> OxmlElement:
    paragraph = OxmlElement("w:p")
    paragraph_properties = OxmlElement("w:pPr")
    paragraph_properties.append(section_properties(landscape))
    paragraph.append(paragraph_properties)
    return paragraph


def format_tables(document) -> list[int]:
    landscape_tables: list[int] = []
    for table_number, table in enumerate(document.tables, 1):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders(table)
        if table.rows:
            set_repeat_table_header(table.rows[0])
        landscape = table_requires_landscape(table)
        target_width = 12960 if landscape else 7920
        apply_table_widths(table, allocate_widths(table, target_width))
        column_count = len(table.columns)
        font_size = 7.5 if column_count >= 7 else 8 if column_count >= 5 else 8.5 if column_count == 4 else 9
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                if row_index == 0:
                    set_cell_shading(cell, "D9D9D9")
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.widow_control = True
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(font_size)
                        if row_index == 0:
                            run.font.bold = True
        caption_node = table._tbl.getprevious()
        while caption_node is not None and caption_node.tag == qn("w:p"):
            caption = Paragraph(caption_node, table._parent)
            if caption.text.strip():
                break
            caption_node = caption_node.getprevious()
        else:
            caption = None
        if caption is None or not caption.text.strip().startswith("Table "):
            raise RuntimeError(f"Table {table_number} has no caption immediately above it")
        caption.style = document.styles["Caption"]
        caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
        caption.paragraph_format.keep_with_next = True
        if landscape:
            landscape_tables.append(table_number)
    return landscape_tables


def scale_figures(document) -> None:
    maximum_width = Inches(5.9)
    maximum_height = Inches(7.25)
    for shape in document.inline_shapes:
        width_ratio = maximum_width / shape.width
        height_ratio = maximum_height / shape.height
        scale = min(1.0, width_ratio, height_ratio)
        if scale < 1.0:
            shape.width = int(shape.width * scale)
            shape.height = int(shape.height * scale)


def style_imported_body(document) -> tuple[int, int, int]:
    figure_caption_count = 0
    table_caption_count = 0
    image_paragraph_count = 0
    references_active = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name in {"Contents 1", "Contents 2", "List Entry"}:
            continue
        if text == "References":
            references_active = True
        elif text == "Appendices":
            references_active = False

        if paragraph.style.name == "Heading 1":
            paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.keep_with_next = True
        elif paragraph.style.name in {"Heading 2", "Heading 3"}:
            paragraph.paragraph_format.keep_with_next = True

        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.keep_with_next = True
            image_paragraph_count += 1
        elif paragraph._p.xpath(".//m:oMathPara"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(4)
        elif paragraph.style.name == "Normal" and text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.widow_control = True

        caption_match = re.match(
            r"^(Figure|Table)\s+((?:\d+(?:\.\d+)*)|(?:[A-Z]\.\d+))\.\s+",
            text,
        )
        if caption_match:
            kind = caption_match.group(1)
            paragraph.style = document.styles["Caption"]
            if kind == "Figure":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                append_tc_field(paragraph, text, "F")
                figure_caption_count += 1
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                append_tc_field(paragraph, text, "T")
                table_caption_count += 1

        if references_active and text and paragraph.style.name == "Normal":
            paragraph.style = document.styles["Reference"]
    return figure_caption_count, table_caption_count, image_paragraph_count


def set_update_fields(document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def force_new_page_sections(document) -> None:
    for section in document.sections:
        section_type = section._sectPr.find(qn("w:type"))
        if section_type is None:
            section_type = OxmlElement("w:type")
            section._sectPr.insert(0, section_type)
        section_type.set(qn("w:val"), "nextPage")


def load_page_map() -> dict[str, str]:
    if not PAGINATION_MAP.is_file():
        return {}
    data = json.loads(PAGINATION_MAP.read_text(encoding="utf-8"))
    return {str(key): str(value) for key, value in data.get("display_pages", {}).items()}


template_document = Document(TEMPLATE)
template_texts = [paragraph.text for paragraph in template_document.paragraphs]
raw_document = Document(RAW_BODY)

raw_elements = [
    element
    for element in raw_document._element.body.iterchildren()
    if element.tag != qn("w:sectPr")
]
abstract_start = next(
    index for index, element in enumerate(raw_elements) if direct_child_text(element) == "Abstract"
)
chapter_start = next(
    index
    for index, element in enumerate(raw_elements)
    if direct_child_text(element).startswith("Chapter 1 Introduction")
)
abstract_elements = raw_elements[abstract_start:chapter_start]
body_elements = reorder_table_captions(raw_elements[chapter_start:])

heading_entries: list[tuple[str, int]] = [
    ("Abstract", 1),
    ("Acknowledgements", 1),
    ("List of Figures", 1),
    ("List of Tables", 1),
]
for element in body_elements:
    text = direct_child_text(element)
    style_id = paragraph_style_id(element)
    if text and style_id in {"Heading1", "Heading2", "Heading3"}:
        level = {"Heading1": 1, "Heading2": 2, "Heading3": 3}[style_id]
        heading_entries.append((text, level))

figure_entries = [
    direct_child_text(element)
    for element in body_elements
    if re.match(r"^Figure\s+(?:\d+(?:\.\d+)+|[A-Z]\.\d+)\.\s+", direct_child_text(element))
]
table_entries = [
    direct_child_text(element)
    for element in body_elements
    if re.match(r"^Table\s+(?:\d+(?:\.\d+)+|[A-Z]\.\d+)\.\s+", direct_child_text(element))
]

document = Document()
clear_body(document)
configure_styles(document)
set_section_geometry(document.sections[0])
document.sections[0].start_type = WD_SECTION.NEW_PAGE
document.sections[0].footer.is_linked_to_previous = False
clear_paragraph(document.sections[0].footer.paragraphs[0])
copy_numbering(raw_document, document)

document.core_properties.title = TITLE
document.core_properties.author = ""
document.core_properties.subject = "NTU dissertation draft assembled from the official template structure"
document.core_properties.keywords = "knowledge graph recommender systems; explainability; dissertation"

add_cover_and_forms(document, template_texts)

front_section = document.add_section(WD_SECTION.NEW_PAGE)
set_section_geometry(front_section)
configure_numbered_footer(front_section, "lowerRoman", 1)

page_map = load_page_map()
add_static_index(document, "Table of Contents", heading_entries, page_map)
document.add_page_break()

abstract_wrappers = [
    append_imported_element(element, raw_document, document) for element in abstract_elements
]
abstract_heading = next(
    wrapper for wrapper in abstract_wrappers if isinstance(wrapper, Paragraph) and wrapper.text.strip() == "Abstract"
)
abstract_heading.style = document.styles["Front Matter Heading"]
abstract_heading.paragraph_format.page_break_before = False

document.add_page_break()
document.add_paragraph("Acknowledgements", style="Front Matter Heading")
placeholder = document.add_paragraph("[Acknowledgements to be filled.]")
placeholder.alignment = WD_ALIGN_PARAGRAPH.LEFT

document.add_page_break()
add_list_index(document, "List of Figures", figure_entries, page_map)
document.add_page_break()
add_list_index(document, "List of Tables", table_entries, page_map)

body_section = document.add_section(WD_SECTION.NEW_PAGE)
set_section_geometry(body_section)
configure_numbered_footer(body_section, "decimal", 1)

element_index = 0
while element_index < len(body_elements):
    element = body_elements[element_index]
    is_landscape_pair = False
    if (
        element.tag == qn("w:p")
        and direct_child_text(element).startswith("Table ")
        and element_index + 1 < len(body_elements)
        and body_elements[element_index + 1].tag == qn("w:tbl")
    ):
        source_table = Table(body_elements[element_index + 1], raw_document._body)
        is_landscape_pair = table_requires_landscape(source_table)
    if is_landscape_pair:
        landscape_section = document.add_section(WD_SECTION.NEW_PAGE)
        set_section_geometry(landscape_section, landscape=True)
        continue_page_numbering(landscape_section)
        append_imported_element(element, raw_document, document)
        append_rebuilt_table(
            Table(body_elements[element_index + 1], raw_document._body),
            document,
        )
        portrait_section = document.add_section(WD_SECTION.NEW_PAGE)
        set_section_geometry(portrait_section, landscape=False)
        continue_page_numbering(portrait_section)
        element_index += 2
        continue
    if element.tag == qn("w:tbl"):
        append_rebuilt_table(Table(element, raw_document._body), document)
    else:
        append_imported_element(element, raw_document, document)
    element_index += 1

figure_caption_count, table_caption_count, image_paragraph_count = style_imported_body(document)
landscape_table_numbers = format_tables(document)
scale_figures(document)
set_update_fields(document)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
AUDIT.parent.mkdir(parents=True, exist_ok=True)
document.save(OUTPUT)

reopened = Document(OUTPUT)
paragraph_texts = [paragraph.text.strip() for paragraph in reopened.paragraphs]
heading_one = [
    paragraph.text.strip()
    for paragraph in reopened.paragraphs
    if paragraph.style.name == "Heading 1"
]
reference_start = next(
    index
    for index, paragraph in enumerate(reopened.paragraphs)
    if paragraph.text.strip() == "References" and paragraph.style.name == "Heading 1"
)
appendix_start = next(
    index
    for index, paragraph in enumerate(reopened.paragraphs)
    if paragraph.text.strip() == "Appendices" and paragraph.style.name == "Heading 1"
)
reference_entries = [
    paragraph.text.strip()
    for paragraph in reopened.paragraphs[reference_start + 1 : appendix_start]
    if paragraph.text.strip() and paragraph.style.name == "Reference"
]

AUDIT.write_text(
    json.dumps(
        {
            "template": str(TEMPLATE.relative_to(ROOT)),
            "template_sha256": digest(TEMPLATE),
            "raw_body": str(RAW_BODY.relative_to(ROOT)),
            "raw_body_sha256": digest(RAW_BODY),
            "output": str(OUTPUT.relative_to(ROOT)),
            "output_sha256": digest(OUTPUT),
            "template_used_as_mother_document": False,
            "template_used_as_layout_form_and_asset_source": True,
            "template_package_sanitized_to_remove_sample_section_artifacts": True,
            "paragraphs": len(reopened.paragraphs),
            "tables": len(reopened.tables),
            "inline_shapes": len(reopened.inline_shapes),
            "sections": len(reopened.sections),
            "landscape_sections": sum(
                section.orientation == WD_ORIENT.LANDSCAPE for section in reopened.sections
            ),
            "landscape_table_numbers": landscape_table_numbers,
            "chapter_heading_count": sum(
                bool(re.match(r"^Chapter [1-6]\b", text)) for text in heading_one
            ),
            "reference_entries": len(reference_entries),
            "figure_captions": figure_caption_count,
            "table_captions": table_caption_count,
            "image_paragraphs": image_paragraph_count,
            "static_toc_entries": len(heading_entries),
            "static_list_of_figures_entries": len(figure_entries),
            "static_list_of_tables_entries": len(table_entries),
            "pagination_map_applied": bool(page_map),
            "manual_placeholders_preserved": all(
                marker in "\n".join(paragraph_texts)
                for marker in [AUTHOR, SUPERVISOR, DATE]
            ),
        },
        indent=2,
    )
    + "\n",
    encoding="utf-8",
)
print(OUTPUT)
