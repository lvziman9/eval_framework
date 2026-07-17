from __future__ import annotations

from collections import Counter
from hashlib import sha256
import json
from pathlib import Path
import re
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[3]
SOURCE = ROOT / "paper/current_dissertation/FULL_DISSERTATION_CURRENT.md"
BIBLIOGRAPHY = ROOT / "paper/current_dissertation/references/REFERENCES_CURRENT.bib"
TEMPLATE = ROOT / "paper/templates/NTU_template.docx"
RAW_BODY = ROOT / "paper/docx_build_v2/intermediate/DISSERTATION_BODY_PANDOC_RAW.docx"
DOCX = ROOT / "paper/docx_build_v2/DISSERTATION_NTU_TEMPLATE_DRAFT_V2.docx"
ASSEMBLY_AUDIT = ROOT / "paper/docx_build_v2/reports/ASSEMBLY_AUDIT.json"
RENDER_AUDIT = ROOT / "paper/docx_build_v2/reports/RENDER_AUDIT_FINAL.json"
PAGINATION_MAP = ROOT / "paper/docx_build_v2/intermediate/PAGINATION_MAP.json"
OUTPUT = ROOT / "paper/docx_build_v2/reports/FINAL_AUDIT.json"

EXPECTED_HASHES = {
    "source": "90e782afa0d7d400bd45c7b80cda6fbe5a0ac67e1dfd79ccdf1b514df1dadb7c",
    "bibliography": "3c05b30da758ea8f60ceef4f68c9f73753818f90057d3955ec113fd6a02baf82",
    "template": "cd82dc0a40d8a8a6e3b093b144c8f3573578c83ed3f2990913a94d8ca2cdaff4",
}


def digest(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def source_citation_keys(text: str) -> set[str]:
    keys: set[str] = set()
    for group in re.findall(r"\[([^\]]*@[a-z][A-Za-z0-9_:.+\-/]+[^\]]*)\]", text):
        keys.update(re.findall(r"@([a-z][A-Za-z0-9_:.+\-/]+)", group))
    return keys


def decimal_tokens(text: str) -> Counter[str]:
    return Counter(re.findall(r"(?<![A-Za-z])[-+]?\d+\.\d+(?:[eE][-+]?\d+)?", text))


source_text = SOURCE.read_text(encoding="utf-8")
source_keys = source_citation_keys(source_text)
document = Document(DOCX)
raw_document = Document(RAW_BODY)
render_audit = json.loads(RENDER_AUDIT.read_text(encoding="utf-8"))
assembly_audit = json.loads(ASSEMBLY_AUDIT.read_text(encoding="utf-8"))
pagination_map = json.loads(PAGINATION_MAP.read_text(encoding="utf-8"))

paragraph_texts = [paragraph.text.strip() for paragraph in document.paragraphs]
all_text = "\n".join(paragraph_texts)
heading_one = [
    paragraph.text.strip()
    for paragraph in document.paragraphs
    if paragraph.style.name == "Heading 1"
]
reference_start = next(
    index
    for index, paragraph in enumerate(document.paragraphs)
    if paragraph.text.strip() == "References" and paragraph.style.name == "Heading 1"
)
appendix_start = next(
    index
    for index, paragraph in enumerate(document.paragraphs)
    if paragraph.text.strip() == "Appendices" and paragraph.style.name == "Heading 1"
)
reference_entries = [
    paragraph.text.strip()
    for paragraph in document.paragraphs[reference_start + 1 : appendix_start]
    if paragraph.text.strip() and paragraph.style.name == "Reference"
]

source_table_cells = [
    cell.text
    for table in raw_document.tables
    for row in table.rows
    for cell in row.cells
]
output_table_cells = [
    cell.text
    for table in document.tables
    for row in table.rows
    for cell in row.cells
]

source_figure_paths = sorted((ROOT / "paper/current_dissertation/figures/png").glob("*.png"))
source_figure_hashes = {digest(path): path.name for path in source_figure_paths}
logo = ROOT / "paper/docx_build_v2/intermediate/template_media/image1.jpg"

with ZipFile(DOCX) as archive:
    zip_integrity_error = archive.testzip()
    media = {
        sha256(archive.read(name)).hexdigest(): name
        for name in archive.namelist()
        if name.startswith("word/media/")
    }
    document_xml = archive.read("word/document.xml").decode("utf-8")
    styles_xml = archive.read("word/styles.xml").decode("utf-8")
    footer_xml = "\n".join(
        archive.read(name).decode("utf-8")
        for name in archive.namelist()
        if re.match(r"word/footer\d+\.xml$", name)
    )

contents_paragraphs = [
    paragraph
    for paragraph in document.paragraphs
    if paragraph.style.name in {"Contents 1", "Contents 2"}
]
list_paragraphs = [
    paragraph for paragraph in document.paragraphs if paragraph.style.name == "List Entry"
]

section_page_number_formats = []
for section in document.sections:
    page_number_type = section._sectPr.find(qn("w:pgNumType"))
    if page_number_type is not None:
        section_page_number_formats.append(
            {
                "format": page_number_type.get(qn("w:fmt")),
                "start": page_number_type.get(qn("w:start")),
            }
        )

source_decimals = decimal_tokens("\n".join(source_table_cells))
output_decimals = decimal_tokens("\n".join(output_table_cells))

checks = {
    "source_markdown_unchanged": digest(SOURCE) == EXPECTED_HASHES["source"],
    "bibliography_unchanged": digest(BIBLIOGRAPHY) == EXPECTED_HASHES["bibliography"],
    "template_unchanged": digest(TEMPLATE) == EXPECTED_HASHES["template"],
    "docx_exists": DOCX.is_file(),
    "docx_sha256": digest(DOCX),
    "zip_integrity": zip_integrity_error is None,
    "paragraphs": len(document.paragraphs),
    "tables": len(document.tables),
    "inline_shapes_total": len(document.inline_shapes),
    "body_figure_count": len(document.inline_shapes) - 1,
    "sections": len(document.sections),
    "landscape_sections": sum(
        section.orientation == WD_ORIENT.LANDSCAPE for section in document.sections
    ),
    "chapter_heading_count": sum(
        bool(re.match(r"^Chapter [1-6]\b", text)) for text in heading_one
    ),
    "references_present": "References" in heading_one,
    "appendices_present": "Appendices" in heading_one,
    "reference_entries": len(reference_entries),
    "source_citation_key_count": len(source_keys),
    "unresolved_citation_syntax_absent": not bool(re.search(r"\[@[a-z]", all_text)),
    "table_cell_text_exact": source_table_cells == output_table_cells,
    "table_decimal_tokens_preserved": source_decimals == output_decimals,
    "table_decimal_token_count": sum(source_decimals.values()),
    "figure_caption_count": sum(
        paragraph.style.name == "Caption" and paragraph.text.strip().startswith("Figure ")
        for paragraph in document.paragraphs
    ),
    "table_caption_count": sum(
        paragraph.style.name == "Caption" and paragraph.text.strip().startswith("Table ")
        for paragraph in document.paragraphs
    ),
    "source_figure_count": len(source_figure_paths),
    "source_figures_embedded_unchanged": all(hash_value in media for hash_value in source_figure_hashes),
    "missing_source_figures": [
        name for hash_value, name in source_figure_hashes.items() if hash_value not in media
    ],
    "template_logo_embedded_unchanged": digest(logo) in media,
    "media_part_count": len(media),
    "hyperlink_relationship_count": sum(
        "hyperlink" in relationship.reltype for relationship in document.part.rels.values()
    ),
    "toc_entries": len(contents_paragraphs),
    "list_entries": len(list_paragraphs),
    "toc_page_numbers_complete": all(
        paragraph.text.split("\t")[-1].strip() for paragraph in contents_paragraphs
    ),
    "list_page_numbers_complete": all(
        paragraph.text.split("\t")[-1].strip() for paragraph in list_paragraphs
    ),
    "pagination_map_applied": assembly_audit["pagination_map_applied"],
    "pagination_map_heading_count": pagination_map["mapped_heading_count"],
    "pagination_map_figure_count": pagination_map["mapped_figure_count"],
    "pagination_map_table_count": pagination_map["mapped_table_count"],
    "page_number_formats": section_page_number_formats,
    "page_field_count": footer_xml.count(" PAGE "),
    "figure_tc_fields": document_xml.count(' TC "Figure ') + document_xml.count(" TC &quot;Figure "),
    "table_tc_fields": document_xml.count(' TC "Table ') + document_xml.count(" TC &quot;Table "),
    "caption_style_explicit_black": 'w:styleId="Caption"' in styles_xml and 'w:color w:val="000000"' in styles_xml,
    "render_page_count": render_audit["page_count"],
    "render_portrait_pages": render_audit["portrait_pages"],
    "render_landscape_pages": render_audit["landscape_pages"],
    "render_blank_pages": render_audit["blank_pages"],
    "render_edge_touch_pages": render_audit["edge_touch_pages"],
    "render_replacement_character_pages": render_audit["replacement_character_pages"],
    "render_image_error_marker_pages": render_audit["image_error_marker_pages"],
}

required_boolean_checks = [
    "source_markdown_unchanged",
    "bibliography_unchanged",
    "template_unchanged",
    "docx_exists",
    "zip_integrity",
    "references_present",
    "appendices_present",
    "unresolved_citation_syntax_absent",
    "table_cell_text_exact",
    "table_decimal_tokens_preserved",
    "source_figures_embedded_unchanged",
    "template_logo_embedded_unchanged",
    "toc_page_numbers_complete",
    "list_page_numbers_complete",
    "pagination_map_applied",
    "caption_style_explicit_black",
]
failures = [key for key in required_boolean_checks if not checks[key]]
expected_counts = {
    "tables": 16,
    "inline_shapes_total": 12,
    "body_figure_count": 11,
    "landscape_sections": 14,
    "chapter_heading_count": 6,
    "reference_entries": 29,
    "source_citation_key_count": 29,
    "figure_caption_count": 11,
    "table_caption_count": 16,
    "source_figure_count": 11,
    "media_part_count": 12,
    "toc_entries": 49,
    "list_entries": 27,
    "pagination_map_heading_count": 49,
    "pagination_map_figure_count": 11,
    "pagination_map_table_count": 16,
    "page_field_count": 2,
    "figure_tc_fields": 11,
    "table_tc_fields": 16,
    "render_page_count": 92,
    "render_portrait_pages": 77,
    "render_landscape_pages": 15,
}
for key, expected in expected_counts.items():
    if checks[key] != expected:
        failures.append(key)
for key in [
    "render_blank_pages",
    "render_edge_touch_pages",
    "render_replacement_character_pages",
    "render_image_error_marker_pages",
    "missing_source_figures",
]:
    if checks[key]:
        failures.append(key)
if checks["hyperlink_relationship_count"] < 20:
    failures.append("hyperlink_relationship_count")
if checks["page_number_formats"] != [
    {"format": "lowerRoman", "start": "1"},
    {"format": "decimal", "start": "1"},
]:
    failures.append("page_number_formats")

checks["audit_failures"] = failures
checks["audit_result"] = "PASS" if not failures else "FAIL"
OUTPUT.write_text(json.dumps(checks, indent=2, default=dict) + "\n", encoding="utf-8")
print(json.dumps(checks, indent=2, default=dict))
if failures:
    raise SystemExit(1)
